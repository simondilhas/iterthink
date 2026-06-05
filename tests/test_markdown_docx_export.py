"""Tests for Word export (markdown → docx)."""

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import pytest
from docx import Document

from iterthink.services import docx_placeholders, markdown_docx_export
from iterthink.services.markdown_docx_export import BLOCK_GAP_PT, ExportMeta

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_BLOCK_GAP_TWIPS = str(BLOCK_GAP_PT * 20)


def _minimal_template(tmp_path: Path) -> Path:
    p = tmp_path / "tpl.docx"
    d = Document()
    d.add_paragraph("Title: {Titel} on {Date} by {Author}")
    d.save(p)
    return p


def test_list_docx_templates_includes_bundled_when_present(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Avoid writing under real Documents/.iterthink (blocked in sandbox CI)."""
    from iterthink import config

    monkeypatch.setattr(config, "STORE_DIR", tmp_path / "iterthink_store")
    names = [lbl for lbl, _ in markdown_docx_export.list_docx_templates()]
    assert "Iterthink Standard" in names


def test_apply_docx_placeholders_replaces_in_body(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    doc = Document(str(tpl))
    docx_placeholders.apply_docx_placeholders(
        doc,
        {"{Titel}": "X", "{Date}": "2099-01-01", "{Author}": "Y", "{Name}": "Y"},
    )
    out = tmp_path / "out.docx"
    doc.save(out)
    assert "{Titel}" not in _document_xml_text(out)
    assert "Title: X on 2099-01-01 by Y" in _document_xml_text(out)


def test_markdown_to_docx_inserts_line_breaks_for_single_newlines(tmp_path: Path) -> None:
    """Single newlines inside a paragraph become w:br (Word), not collapsed spaces."""
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("first line\nsecond line\n\nnew paragraph", encoding="utf-8")
    out = tmp_path / "out.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "first line" in body
    assert "second line" in body
    assert "new paragraph" in body
    assert "<w:br" in body


def test_markdown_to_docx_placeholders_include_name(tmp_path: Path) -> None:
    tpl = tmp_path / "tpl.docx"
    d = Document()
    d.add_paragraph("By {Name} on {Date}")
    d.save(tpl)
    md = tmp_path / "note.md"
    md.write_text("Hi", encoding="utf-8")
    out = tmp_path / "out.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="Pat", date_iso="2099-03-03"),
    )
    body = _document_xml_text(out)
    assert "{Name}" not in body
    assert "Pat" in body


def _document_xml_text(docx: Path) -> str:
    with zipfile.ZipFile(docx) as z:
        return z.read("word/document.xml").decode("utf-8")


def _paragraph_spacing_for_text(body_xml: str, needle: str) -> tuple[str | None, str | None]:
    """Return ``(space_after, space_before)`` twips for the first ``w:p`` containing ``needle``."""
    root = ET.fromstring(body_xml)
    for p in root.iter(f"{{{_W_NS}}}p"):
        texts = "".join(t.text or "" for t in p.iter(f"{{{_W_NS}}}t"))
        if needle not in texts:
            continue
        sp = p.find(f".//{{{_W_NS}}}spacing")
        if sp is None:
            return None, None
        return sp.get(f"{{{_W_NS}}}after"), sp.get(f"{{{_W_NS}}}before")
    return None, None


def _spacing_before_between_text(body_xml: str, after_needle: str, before_needle: str) -> str | None:
    """Return ``w:before`` twips on a spacer ``w:p`` between two text needles."""
    start = body_xml.find(after_needle)
    end = body_xml.find(before_needle, start + len(after_needle) if start >= 0 else 0)
    if start < 0 or end < 0 or end <= start:
        return None
    segment = body_xml[start:end]
    matches = re.findall(r'w:before="(\d+)"', segment)
    return matches[-1] if matches else None


def test_markdown_to_docx_smoke(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("# H\n\nHello [^a]\n\n[^a]: Foot\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    assert out.is_file()
    body = _document_xml_text(out)
    assert "2099-02-02" in body
    with zipfile.ZipFile(out) as z:
        names = z.namelist()
    assert "word/footnotes.xml" in names


def test_markdown_to_docx_list_markers_without_template_list_styles(tmp_path: Path) -> None:
    """Bundled template uses Word numbering (no literal U+2022 in document.xml); content must export."""
    tpl = markdown_docx_export.bundled_templates_dir() / "Iterthink Standard.docx"
    if not tpl.is_file():
        pytest.skip("Bundled DOCX template missing")
    md = tmp_path / "note.md"
    md.write_text("- Alpha\n- Beta\n\n1. One\n2. Two\n", encoding="utf-8")
    out = tmp_path / "lists.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "w:numPr" in body
    assert "Alpha" in body
    assert "Beta" in body
    assert "One" in body
    assert "Two" in body


def test_markdown_to_docx_nested_bullet_list_ilvl(tmp_path: Path) -> None:
    """Nested markdown list items must get distinct Word list levels (w:ilvl)."""
    tpl = markdown_docx_export.bundled_templates_dir() / "Iterthink Standard.docx"
    if not tpl.is_file():
        pytest.skip("Bundled DOCX template missing")
    md = tmp_path / "note.md"
    md.write_text(
        "- L1\n"
        "  - L2\n"
        "    - L3\n",
        encoding="utf-8",
    )
    out = tmp_path / "nested.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "L1" in body and "L2" in body and "L3" in body
    assert body.count('<w:ilvl w:val="0"/>') >= 1
    assert body.count('<w:ilvl w:val="1"/>') >= 1
    assert body.count('<w:ilvl w:val="2"/>') >= 1


def test_markdown_to_docx_nested_list_siblings_same_depth(tmp_path: Path) -> None:
    """Sibling items under one parent must stay nested (same w:ilvl), not promoted to depth 0."""
    tpl = markdown_docx_export.bundled_templates_dir() / "Iterthink Standard.docx"
    if not tpl.is_file():
        pytest.skip("Bundled DOCX template missing")
    md = tmp_path / "note.md"
    md.write_text(
        "- L1\n"
        "  - L2a\n"
        "  - L2b\n",
        encoding="utf-8",
    )
    out = tmp_path / "siblings.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "L2a" in body and "L2b" in body
    assert body.count('<w:ilvl w:val="1"/>') >= 2


def test_markdown_to_docx_paragraph_gap_between_body_paragraphs(tmp_path: Path) -> None:
    """Double-Enter between body paragraphs must yield visible gap (space_after on first block)."""
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("Paragraph one.\n\nParagraph two.\n", encoding="utf-8")
    out = tmp_path / "gap.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    after_one, _ = _paragraph_spacing_for_text(body, "Paragraph one.")
    after_two, before_two = _paragraph_spacing_for_text(body, "Paragraph two.")
    assert after_one == _BLOCK_GAP_TWIPS
    assert before_two == _BLOCK_GAP_TWIPS
    assert after_two == _BLOCK_GAP_TWIPS


def test_markdown_to_docx_paragraph_before_table_has_gap(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text(
        "Before table.\n\n| H |\n|---|\n| x |\n\nAfter table.\n",
        encoding="utf-8",
    )
    out = tmp_path / "table_gap.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    after_before, _ = _paragraph_spacing_for_text(body, "Before table.")
    _, before_after = _paragraph_spacing_for_text(body, "After table.")
    assert after_before == _BLOCK_GAP_TWIPS
    assert before_after == _BLOCK_GAP_TWIPS


def test_markdown_to_docx_gap_between_two_tables(tmp_path: Path) -> None:
    """Blank line between GFM tables must yield visible gap in Word (spacer paragraph)."""
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text(
        "| A |\n|---|\n| 1 |\n\n| B |\n|---|\n| 2 |\n",
        encoding="utf-8",
    )
    out = tmp_path / "two_tables.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "A" in body and "B" in body
    gap_before = _spacing_before_between_text(body, ">1<", ">B<")
    assert gap_before == _BLOCK_GAP_TWIPS


def test_markdown_to_docx_exports_gfm_table(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text(
        "| H1 | H2 |\n|---|---|\n| a | bb |\n\nAfter table.\n",
        encoding="utf-8",
    )
    out = tmp_path / "table.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "<w:tbl" in body
    assert "H1" in body
    assert "H2" in body
    assert "bb" in body
    assert "After table." in body


def test_markdown_to_docx_table_inline_bold(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("| Col |\n|-----|\n| **bold** |\n", encoding="utf-8")
    out = tmp_path / "bold_cell.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "<w:tbl" in body
    assert "bold" in body
    assert "<w:b" in body or "w:val=\"true\"" in body


def test_markdown_to_docx_hr_inserts_page_break(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("Before\n\n---\n\nAfter", encoding="utf-8")
    out = tmp_path / "hr.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="A", date_iso="2099-02-02"),
    )
    body = _document_xml_text(out)
    assert "Before" in body
    assert "After" in body
    assert 'w:br w:type="page"' in body
    assert "— — —" not in body


def test_markdown_to_docx_paragraph_comment(tmp_path: Path) -> None:
    tpl = _minimal_template(tmp_path)
    md = tmp_path / "note.md"
    md.write_text("# Title\n\nFirst body.\n\nSecond body.\n", encoding="utf-8")
    out = tmp_path / "out_comments.docx"
    markdown_docx_export.markdown_to_docx(
        markdown_src=md.read_text(encoding="utf-8"),
        md_path=md,
        template_path=tpl,
        output_path=out,
        meta=ExportMeta(title_stem="note", author="Author", date_iso="2099-01-01", comment_author="Author"),
        paragraph_comments={1: "Annotation on first body paragraph."},
    )
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        assert "word/comments.xml" in z.namelist()
        cxml = z.read("word/comments.xml").decode("utf-8")
    assert "Annotation on first body paragraph" in cxml
