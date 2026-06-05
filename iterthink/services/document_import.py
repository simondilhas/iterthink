"""Convert DOCX / PDF imports to Markdown with assets beside the target document."""

from __future__ import annotations

import hashlib
import logging
import math
import os
import re
import statistics
from collections import defaultdict
from pathlib import Path
from typing import Literal

from iterthink import config
from iterthink.ocr_settings import IMAGE_IMPORT_EXTENSIONS

logger = logging.getLogger(__name__)

PdfProfileHeuristic = Literal["text", "plan"]

PDF_RENDER_SCALE_TEXT = 1.5
PDF_RENDER_SCALE_PLAN = 2.0

ALLOWED_IMPORT_EXTENSIONS = frozenset({"docx", "pdf", *IMAGE_IMPORT_EXTENSIONS})

# Trailing ``_<long digits>`` from some OS file pickers (temp copy ids).
_PICKER_TEMP_STEM_SUFFIX = re.compile(r"_\d{10,}$")


def normalize_import_library_stem(stem: str) -> str:
    """Sanitize and drop picker temp numeric suffixes from the library note stem."""
    safe = "".join(c for c in stem if c.isalnum() or c in " ._-")[:200].strip()
    if not safe:
        return ""
    m = _PICKER_TEMP_STEM_SUFFIX.search(safe)
    if m and m.start() > 0:
        trimmed = safe[: m.start()].rstrip(" ._-")
        if trimmed:
            return trimmed
    return safe


def import_dest_md_path(src: Path, base: Path) -> Path | None:
    """Target ``.md`` path for an import, or None if the derived name is invalid."""
    safe = normalize_import_library_stem(src.stem)
    if not safe:
        return None
    return base.resolve() / f"{safe}.md"


def import_target_display_path(dest_md: Path, documents_root: Path) -> str:
    """Human path like ``projects/foo/Name.pdf`` (tree-style label, not on-disk ``.md``)."""
    dest_md = dest_md.resolve()
    root = documents_root.resolve()
    try:
        rel_parent = dest_md.parent.relative_to(root)
        folder = "." if rel_parent == Path(".") else rel_parent.as_posix()
    except ValueError:
        folder = dest_md.parent.as_posix()
    return f"{folder}/{dest_md.stem}.pdf"


def import_pdf_dialog_hint(
    dest_md: Path, documents_root: Path, *, import_into_existing: bool
) -> str:
    target = import_target_display_path(dest_md, documents_root)
    if import_into_existing:
        return f"Add version to: {target}"
    return f"Save as: {target}"


def validate_extension(path: Path) -> str | None:
    """Return lower-case extension if allowed, else None."""
    suf = path.suffix.lower().lstrip(".")
    return suf if suf in ALLOWED_IMPORT_EXTENSIONS else None


def docx_to_markdown(src: Path, asset_dir: Path) -> str:
    """Extract markdown from Word; images written under ``asset_dir`` with relative links."""
    from docx import Document as DocxDocument

    asset_dir.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument(str(src))
    markdown_lines: list[str] = []
    stem_safe = re.sub(r"[^\w\-]+", "_", src.stem)[:80] or "doc"
    image_counter = 0

    def _blank_line_before_heading() -> None:
        """Markdown needs a blank line before a heading when following body text."""
        if markdown_lines and markdown_lines[-1] != "":
            markdown_lines.append("")

    for para in doc.paragraphs:
        text = para.text.strip()
        name = para.style.name if para.style else ""
        if not text:
            markdown_lines.append("")
            continue
        if name.startswith("Heading 1"):
            _blank_line_before_heading()
            markdown_lines.append(f"# {text}")
        elif name.startswith("Heading 2"):
            _blank_line_before_heading()
            markdown_lines.append(f"## {text}")
        elif name.startswith("Heading 3"):
            _blank_line_before_heading()
            markdown_lines.append(f"### {text}")
        else:
            markdown_lines.append(text)

    for shape in doc.inline_shapes:
        try:
            image_part = shape._inline.graphic.graphicData.pic.blip.fill.blip
            r_id = image_part.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_id:
                continue
            related_part = doc.part.related_parts[r_id]
            image_bytes = related_part.blob
            content_type = getattr(related_part, "content_type", "") or "image/png"
            ext = content_type.split("/")[-1].lower()
            if ext in ("jpeg", "jpg"):
                ext = "jpg"
            elif ext != "png":
                ext = "png"
            image_filename = f"{stem_safe}_image_{image_counter}.{ext}"
            image_path = asset_dir / image_filename
            image_path.write_bytes(image_bytes)
            rel = os.path.relpath(image_path, asset_dir.parent)
            markdown_lines.append(f"\n![Image]({rel})\n")
            image_counter += 1
        except Exception as ex:
            logger.debug("Inline image skipped: %s", ex)

    parts: list[str] = []
    buf: list[str] = []
    for line in markdown_lines:
        if line == "":
            if buf:
                parts.append("\n".join(buf))
                buf = []
            continue
        buf.append(line)
    if buf:
        parts.append("\n".join(buf))

    return "\n\n".join(parts) if parts else ""


# Lines that are only a bullet glyph (text often continues on the next line).
_BULLET_ONLY_LINE = re.compile(r"^[•·▪▸►◦‣⁃]\s*$")
# Inline bullet + content on same line
_INLINE_BULLET = re.compile(r"^[•·▪▸►◦]\s*(.+)$")
_MARKDOWNISH_BULLET = re.compile(r"^[\-\*\+]\s+(.+)$")
_NUMBERED_LINE = re.compile(r"^(\d{1,3})[\.\)]\s+(.*)$")

# Repetition-based strip: same (y_bucket, strip fingerprint) on enough distinct pages.
# Fingerprint folds digit runs to one placeholder so "Footer 5" / "Footer 6" merge; the
# whole PDF line is skipped when its fingerprint matches. Very short non-numeric lines
# are ignored for matching (reduces stripping punctuation-only repeats).
_PDF_STRIP_MIN_PAGE_FRACTION = 0.45
_PDF_STRIP_MIN_TEXT_LEN = 4
_PDF_Y_BUCKET_SCALE = 100
_WS_COLLAPSE = re.compile(r"\s+")
_DIGIT_RUNS = re.compile(r"\d+")
_PAGE_NUMERIC_ONLY = re.compile(r"^\d{1,6}$")


def _normalize_pdf_line_text(s: str) -> str:
    return _WS_COLLAPSE.sub(" ", s.strip())


def _pdf_strip_fingerprint(norm: str) -> str | None:
    """
    Text used in (y, text) strip keys: whitespace-normalized, digit runs → '#'.
    Returns None if the line should not participate in repetition detection.
    """
    folded = _DIGIT_RUNS.sub("#", norm)
    if len(folded) >= _PDF_STRIP_MIN_TEXT_LEN:
        return folded
    if _PAGE_NUMERIC_ONLY.match(norm):
        return "__page_digits__"
    return None


def _pdf_strip_page_threshold(n_pages: int) -> int:
    """Minimum distinct page count for a (y_bucket, text) key to be stripped."""
    if n_pages < 2:
        return 10**9
    return max(2, int(math.ceil(_PDF_STRIP_MIN_PAGE_FRACTION * n_pages)))


def _pdf_line_strip_key(line: dict, page_height: float) -> tuple[int, str] | None:
    """Key for repetition matching: vertical bucket and strip fingerprint (digits folded)."""
    spans = line.get("spans") or []
    if not spans:
        return None
    raw = "".join(sp.get("text", "") for sp in spans)
    norm = _normalize_pdf_line_text(raw)
    fp = _pdf_strip_fingerprint(norm)
    if fp is None:
        return None
    bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
    cy = (float(bbox[1]) + float(bbox[3])) / 2.0
    ph = max(float(page_height), 1e-6)
    y_frac = max(0.0, min(1.0, cy / ph))
    bk = int(round(y_frac * _PDF_Y_BUCKET_SCALE))
    bk = max(0, min(_PDF_Y_BUCKET_SCALE, bk))
    return (bk, fp)


def _plumber_page_to_lines(page: object) -> list[dict]:
    """Build line dicts (``bbox`` + ``spans`` with ``size``) from pdfplumber ``.chars``."""
    # Keep space glyphs: ``(c.get("text") or "").strip()`` drops literal spaces and breaks words.
    chars = [c for c in (page.chars or []) if (c.get("text") or "") != ""]
    if not chars:
        return []
    heights: list[float] = []
    for c in chars:
        t, b = float(c["top"]), float(c["bottom"])
        h = max(b - t, float(c.get("height") or 0.0), 0.001)
        heights.append(h)
    med = statistics.median(heights) if heights else 10.0
    row_unit = max(2.5, med * 0.55)

    buckets: dict[int, list] = defaultdict(list)
    for c in chars:
        cy = (float(c["top"]) + float(c["bottom"])) / 2.0
        key = int(cy / row_unit)
        buckets[key].append(c)

    lines: list[dict] = []
    for _bk in sorted(buckets.keys()):
        row = sorted(buckets[_bk], key=lambda c: float(c["x0"]))
        spans: list[dict] = []
        cur_parts: list[str] = []
        cur_size = 0.0
        cur_font = ""

        def flush_span() -> None:
            nonlocal cur_parts, cur_size, cur_font
            s = "".join(cur_parts)
            if not s.strip():
                cur_parts = []
                cur_size = 0.0
                cur_font = ""
                return
            spans.append({"text": s, "size": cur_size or 11.0})
            cur_parts = []
            cur_size = 0.0
            cur_font = ""

        widths = [
            max(float(c["x1"]) - float(c["x0"]), 0.001)
            for c in row
            if (t := str(c.get("text") or "")) and not t.isspace()
        ]
        char_mw = statistics.median(widths) if widths else max(med * 0.45, 3.0)
        gap_threshold = max(1.2, med * 0.10, char_mw * 0.48)
        last_x1: float | None = None

        for c in row:
            ch = str(c.get("text") or "")
            sz = float(c.get("size") or 0.0)
            fn = str(c.get("fontname") or "")
            x0, x1 = float(c["x0"]), float(c["x1"])
            if cur_parts and (
                cur_font != fn or (cur_size > 0 and sz > 0 and abs(cur_size - sz) > 0.75)
            ):
                flush_span()
            if ch.isspace():
                if not cur_parts or cur_parts[-1] != " ":
                    cur_parts.append(" ")
                if sz > 0.0:
                    cur_size = sz
                cur_font = fn
                last_x1 = x1
                continue
            if last_x1 is not None:
                gap = x0 - last_x1
                if gap > gap_threshold and (not cur_parts or cur_parts[-1] != " "):
                    cur_parts.append(" ")
            cur_parts.append(ch)
            if sz > 0.0:
                cur_size = sz
            cur_font = fn
            last_x1 = x1
        flush_span()
        if not spans:
            continue
        x0 = min(float(c["x0"]) for c in row)
        x1 = max(float(c["x1"]) for c in row)
        y0 = min(float(c["top"]) for c in row)
        y1 = max(float(c["bottom"]) for c in row)
        lines.append({"bbox": [x0, y0, x1, y1], "spans": spans})
    return lines


def _pdf_repeated_strip_keys_lines(pages_data: list[tuple[float, list[dict]]]) -> frozenset[tuple[int, str]]:
    """
    Lines whose (relative-Y bucket, strip fingerprint) appear on enough distinct pages
    (running heads / footers). ``pages_data`` is ``(page_height, lines)`` per page index.
    """
    n = len(pages_data)
    need = _pdf_strip_page_threshold(n)
    key_pages: dict[tuple[int, str], set[int]] = defaultdict(set)
    for pi, (ph, page_lines) in enumerate(pages_data):
        for line in page_lines:
            key = _pdf_line_strip_key(line, ph)
            if key is None:
                continue
            key_pages[key].add(pi)
    return frozenset(k for k, pages in key_pages.items() if len(pages) >= need)


def _strip_inline_bullet(text: str) -> str | None:
    m = _INLINE_BULLET.match(text)
    if m:
        return m.group(1).strip()
    m = _MARKDOWNISH_BULLET.match(text)
    if m:
        return m.group(1).strip()
    return None


def _pdf_dict_to_markdown(src: Path) -> str:
    """
    Structured PDF → Markdown via pdfplumber (char geometry → line/span dicts): reading order,
    font-size headings, bullets / numbered lists, paragraph gaps from vertical positions.

    Repeated running heads/footers: lines whose strip fingerprint (digits folded to ``#``)
    and relative vertical bucket match on enough distinct pages are omitted entirely.
    Single-page PDFs never strip.
    """
    import pdfplumber

    with pdfplumber.open(str(src)) as pdf:
        pages_data: list[tuple[float, list[dict]]] = []
        for page in pdf.pages:
            ph = float(page.height) if page.height else 792.0
            pages_data.append((ph, _plumber_page_to_lines(page)))

        strip_keys = _pdf_repeated_strip_keys_lines(pages_data)

        all_sizes: list[float] = []
        for ph, page_lines in pages_data:
            for line in page_lines:
                sk = _pdf_line_strip_key(line, ph)
                if sk is not None and sk in strip_keys:
                    continue
                for sp in line.get("spans", []):
                    z = sp.get("size") or 0
                    if z and float(z) > 0:
                        all_sizes.append(float(z))
        if len(all_sizes) >= 6:
            sorted_sz = sorted(all_sizes)
            body_med = sorted_sz[len(sorted_sz) // 4]
        else:
            body_med = statistics.median(all_sizes) if all_sizes else 11.0
        body_med = max(6.0, min(body_med, 22.0))

        page_chunks: list[str] = []
        for pi, (ph, page_lines) in enumerate(pages_data):
            page_num = pi + 1
            flat: list[tuple[float, float, float, float, dict]] = []
            for line in page_lines:
                sk = _pdf_line_strip_key(line, ph)
                if sk is not None and sk in strip_keys:
                    continue
                bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
                x0, y0, _x1, y1 = bbox[0], bbox[1], bbox[2], bbox[3]
                flat.append((y0, x0, y1, line))
            flat.sort(key=lambda t: (round(t[0], 2), round(t[1], 2)))

            events: list[tuple[str, tuple]] = []
            pending_bullet = False

            for _y0s, _x0s, _y1s, line in flat:
                spans = line.get("spans") or []
                if not spans:
                    continue
                raw = "".join(sp.get("text", "") for sp in spans)
                text = raw.strip()
                if not text:
                    continue

                if _BULLET_ONLY_LINE.match(text):
                    pending_bullet = True
                    continue

                bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
                ly0, ly1 = float(bbox[1]), float(bbox[3])

                if pending_bullet:
                    events.append(("bullet", (text, ly0, ly1)))
                    pending_bullet = False
                    continue

                nm = _NUMBERED_LINE.match(text)
                if nm:
                    events.append(("ol", (nm.group(1), nm.group(2).strip(), ly0, ly1)))
                    continue

                stripped = _strip_inline_bullet(text)
                if stripped is not None:
                    events.append(("bullet", (stripped, ly0, ly1)))
                    continue

                sizes = [float(sp.get("size") or 0) for sp in spans if (sp.get("size") or 0) > 0]
                max_sz = max(sizes) if sizes else body_med
                ratio = max_sz / body_med if body_med > 0 else 1.0

                long_line = len(text) > 180
                strong_heading = ratio >= 1.55

                if not long_line or strong_heading:
                    if ratio >= 1.48:
                        events.append(("h1", (text, ly0, ly1)))
                        continue
                    if ratio >= 1.26:
                        events.append(("h2", (text, ly0, ly1)))
                        continue
                    if ratio >= 1.11:
                        events.append(("h3", (text, ly0, ly1)))
                        continue

                events.append(("body", (text, ly0, ly1)))

            md_body = _pdf_events_to_markdown(events, body_med)
            marker = f"<!-- page:{page_num} -->"
            if md_body.strip():
                page_chunks.append(f"{marker}\n\n{md_body.strip()}")
            else:
                page_chunks.append(marker)

    return "\n\n".join(page_chunks)


def _pdf_events_to_markdown(events: list[tuple[str, tuple]], body_med: float) -> str:
    """Turn classified line events into Markdown (paragraphs, lists, headings)."""
    gap_para = max(4.0, body_med * 0.55)
    gap_join = max(2.5, body_med * 0.38)

    parts: list[str] = []
    cur_para: list[str] = []
    prev_y1: float | None = None
    list_lines: list[str] = []

    def flush_para() -> None:
        nonlocal cur_para, prev_y1
        if cur_para:
            parts.append("\n".join(cur_para))
        cur_para = []
        prev_y1 = None

    def flush_list() -> None:
        nonlocal list_lines
        if list_lines:
            parts.append("\n".join(list_lines))
        list_lines = []

    for ev in events:
        kind, payload = ev
        if kind == "body":
            flush_list()
            text = str(payload[0])
            y0, y1 = float(payload[1]), float(payload[2])
            if prev_y1 is not None and cur_para and (y0 - prev_y1) <= gap_join:
                cur_para[-1] = (cur_para[-1].rstrip() + " " + text.lstrip()).strip()
            elif prev_y1 is not None and cur_para and (y0 - prev_y1) > gap_para:
                flush_para()
                cur_para.append(text)
            else:
                if cur_para:
                    parts.append("\n".join(cur_para))
                    cur_para = []
                cur_para.append(text)
            prev_y1 = y1
            continue

        flush_para()
        if kind == "h1":
            flush_list()
            parts.append(f"# {payload[0]}")
        elif kind == "h2":
            flush_list()
            parts.append(f"## {payload[0]}")
        elif kind == "h3":
            flush_list()
            parts.append(f"### {payload[0]}")
        elif kind == "bullet":
            list_lines.append(f"- {str(payload[0])}")
        elif kind == "ol":
            list_lines.append(f"{payload[0]}. {payload[1]}")

    flush_para()
    flush_list()

    return "\n\n".join(parts)


def _pdf_to_markdown_pypdf_plain(src: Path) -> str:
    """Per-page ``pypdf`` ``extract_text`` with ``<!-- page:N -->`` markers (no repetition strip)."""
    from pypdf import PdfReader

    reader = PdfReader(str(src))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages):
        body = (page.extract_text() or "").strip()
        marker = f"<!-- page:{i + 1} -->"
        chunks.append(f"{marker}\n\n{body}" if body else marker)
    return "\n\n".join(chunks)


def _pdf_to_markdown_legacy(src: Path) -> str:
    """
    Fallback: pdfplumber-derived lines with repetition strip; per-page ``pypdf`` text if empty.

    Inserts ``<!-- page:N -->`` (1-based) before each page's content for scroll mapping.
    """
    import pdfplumber
    from pypdf import PdfReader

    pages_data: list[tuple[float, list[dict]]] = []
    try:
        with pdfplumber.open(str(src)) as pdf:
            for p in pdf.pages:
                ph = float(p.height) if p.height else 792.0
                pages_data.append((ph, _plumber_page_to_lines(p)))
    except Exception as ex:
        logger.debug("pdfplumber legacy prep failed: %s", ex)

    strip_keys = _pdf_repeated_strip_keys_lines(pages_data) if pages_data else frozenset()

    reader = PdfReader(str(src))
    chunks: list[str] = []
    for page_index, page in enumerate(reader.pages):
        page_num = page_index + 1
        if page_index < len(pages_data):
            ph, page_lines = pages_data[page_index]
        else:
            ph, page_lines = 792.0, []
        lines_filtered: list[str] = []
        for line in page_lines:
            bk = _pdf_line_strip_key(line, ph)
            if bk is not None and bk in strip_keys:
                continue
            raw = "".join(str(sp.get("text", "")) for sp in (line.get("spans") or []))
            if raw.strip():
                lines_filtered.append(raw.strip())
        raw2 = "\n".join(lines_filtered)
        paras = [p.strip() for p in raw2.split("\n\n") if p.strip()]
        if not paras and lines_filtered:
            paras = [" ".join(lines_filtered)]
        if not paras:
            raw = (page.extract_text() or "").strip()
            paras = [p.strip() for p in raw.split("\n\n") if p.strip()] if raw else []
        body = "\n\n".join(paras) if paras else ""
        marker = f"<!-- page:{page_num} -->"
        if body.strip():
            chunks.append(f"{marker}\n\n{body}")
        else:
            chunks.append(marker)

    return "\n\n".join(chunks)


def pdf_to_markdown(src: Path) -> str:
    """
    PDF → Markdown with ``<!-- page:N -->`` markers.

    Primary path: structured pdfplumber extraction (reading order, font-size headings,
    bullets / numbered lists, paragraph gaps). Repeated running heads/footers are dropped when
    the same strip fingerprint and relative vertical position recur on enough distinct pages.
    Then plain per-page ``pypdf`` text, then legacy hybrid (plumber lines + strip + ``pypdf``).
    """

    def _has_body(md: str) -> bool:
        raw = re.sub(r"<!--\s*page:\d+\s*-->", "", md)
        return bool(raw.strip())

    try:
        md_struct = _pdf_dict_to_markdown(src)
        if _has_body(md_struct):
            return md_struct
    except Exception as ex:
        logger.warning("Structured PDF extraction failed: %s", ex)

    try:
        md_plain = _pdf_to_markdown_pypdf_plain(src)
        if _has_body(md_plain):
            return md_plain
    except Exception as ex:
        logger.warning("Plain pypdf PDF extraction failed: %s", ex)

    md = _pdf_to_markdown_legacy(src)
    if not _has_body(md):
        if config.OCR_ENABLED:
            from iterthink.services.ocr_import import pdf_to_markdown_ocr

            return pdf_to_markdown_ocr(src)
        if not md.strip():
            return (
                "<!-- PDF text extraction returned no content. This may be a scanned PDF. -->\n\n"
                "*Note: No extractable text; use Original PDF in Compare to view.*"
            )
        return md
    return md


def pdf_render_scale_for_profile(pdf_profile: PdfProfileHeuristic | None) -> float:
    return PDF_RENDER_SCALE_PLAN if pdf_profile == "plan" else PDF_RENDER_SCALE_TEXT


def extract_pdf_pages_geometry(src: Path) -> dict:
    """
    Plan PDFs: per-page lines with PDF-space bboxes (no markdown flattening).
    """
    import pdfplumber

    with pdfplumber.open(str(src)) as pdf:
        pages_data: list[tuple[float, float, list[dict]]] = []
        for page in pdf.pages:
            ph = float(page.height) if page.height else 792.0
            pw = float(page.width) if page.width else 612.0
            pages_data.append((ph, pw, _plumber_page_to_lines(page)))

        strip_keys = _pdf_repeated_strip_keys_lines([(ph, lines) for ph, _pw, lines in pages_data])

        pages_out: list[dict] = []
        for pi, (ph, pw, page_lines) in enumerate(pages_data):
            lines_out: list[dict] = []
            for line in page_lines:
                sk = _pdf_line_strip_key(line, ph)
                if sk is not None and sk in strip_keys:
                    continue
                spans = line.get("spans") or []
                text = "".join(sp.get("text", "") for sp in spans).strip()
                if not text:
                    continue
                bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
                sizes = [float(sp.get("size") or 0) for sp in spans if (sp.get("size") or 0) > 0]
                lines_out.append(
                    {
                        "text": text,
                        "bbox": [float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])],
                        "size": max(sizes) if sizes else 11.0,
                    }
                )
            pages_out.append({"page": pi + 1, "width": pw, "height": ph, "lines": lines_out})
    return {"pages": pages_out}


def import_pdf_with_profile_and_geometry(
    src: Path, pdf_profile: PdfProfileHeuristic
) -> tuple[str, dict | None]:
    """Markdown body and optional plan geometry (one pdfplumber pass for plan imports)."""
    if pdf_profile == "plan":
        from iterthink.services.plan_text_extract import (
            extract_plan_geometry,
            plan_stub_markdown_minimal,
        )

        geometry = extract_plan_geometry(src)
        return plan_stub_markdown_minimal(), geometry
    return pdf_to_markdown(src), None


def import_plan_pdf_fast_stub() -> str:
    """Minimal markdown for plan import before geometry extraction runs in background."""
    from iterthink.services.plan_text_extract import plan_stub_markdown_minimal

    return plan_stub_markdown_minimal()


def import_pdf_with_profile(src: Path, pdf_profile: PdfProfileHeuristic) -> str:
    """Return markdown body for a PDF import using the chosen profile."""
    return import_pdf_with_profile_and_geometry(src, pdf_profile)[0]


def import_pdf_for_profile(src: Path) -> tuple[str, PdfProfileHeuristic]:
    """Return markdown body and profile for a PDF import (auto-classified)."""
    prof = classify_pdf_profile(src)
    return import_pdf_with_profile(src, prof), prof


def classify_pdf_profile(src: Path) -> PdfProfileHeuristic:
    """
    Rough split: text-heavy PDFs → ``text``; sparse extraction → ``plan`` / drawing-like.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(src))
    n = len(reader.pages)
    if n == 0:
        return "plan"
    total_chars = 0
    for page in reader.pages:
        total_chars += len((page.extract_text() or "").strip())
    avg = total_chars / max(n, 1)
    if avg < 120:
        return "plan"
    return "text"


def pdf_render_cache_dir(pdf_abs: Path) -> Path:
    """Stable cache folder under store for rendered PNG pages."""
    st = pdf_abs.stat()
    key_src = f"{pdf_abs.resolve()}:{st.st_mtime_ns}:{st.st_size}".encode()
    key = hashlib.sha256(key_src).hexdigest()[:24]
    d = config.STORE_DIR / "pdf_render_cache" / key
    d.mkdir(parents=True, exist_ok=True)
    return d


def render_pdf_to_png_pages(
    pdf_abs: Path,
    *,
    pdf_profile: PdfProfileHeuristic | None = None,
) -> list[Path]:
    """
    Rasterize each PDF page to PNG under the render cache (pypdfium2 + Pillow).
    Returns ordered list of PNG paths.
    """
    import pypdfium2 as pdfium

    scale = pdf_render_scale_for_profile(pdf_profile)
    cache = pdf_render_cache_dir(pdf_abs)
    marker = cache / ".source"
    src_tag = f"{pdf_abs.resolve()}:{pdf_abs.stat().st_mtime_ns}:scale={scale}"
    if marker.is_file() and marker.read_text(encoding="utf-8") == src_tag:
        existing = sorted(cache.glob("page_*.png"))
        if existing:
            return existing

    for old in cache.glob("page_*.png"):
        old.unlink(missing_ok=True)
    for old in cache.glob("page_*_textov_*.png"):
        old.unlink(missing_ok=True)

    pdf = pdfium.PdfDocument(str(pdf_abs))
    out: list[Path] = []
    try:
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil()
            p = cache / f"page_{i + 1:04d}.png"
            pil_image.save(str(p))
            out.append(p)
    finally:
        pdf.close()

    marker.write_text(src_tag, encoding="utf-8")
    return out


def import_file_to_markdown(src: Path, md_path: Path) -> str:
    """Route by extension. DOCX assets go to ``<md_stem>_assets/`` next to ``md_path``."""
    ext = validate_extension(src)
    if ext is None:
        raise ValueError(f"Unsupported import type: {src.suffix}")

    if ext == "docx":
        asset_dir = config.IMPORT_ASSETS_DIR / md_path.stem
        return docx_to_markdown(src, asset_dir)

    return import_pdf_for_profile(src)[0]

