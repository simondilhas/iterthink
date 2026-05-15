"""Export Markdown to Word (.docx) using markdown-it-py and a custom token walker."""

from __future__ import annotations

import importlib.util
from dataclasses import dataclass
from pathlib import Path
from typing import Any, NamedTuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.styles import BabelFish

from iterthink import config

from . import docx_footnotes, docx_placeholders


def bundled_templates_dir() -> Path:
    spec = importlib.util.find_spec("iterthink")
    if spec is None or spec.origin is None:
        return Path(__file__).resolve().parents[1] / "templates"
    return Path(spec.origin).resolve().parent / "templates"


def user_templates_dir() -> Path:
    d = config.STORE_DIR / "templates"
    d.mkdir(parents=True, exist_ok=True)
    return d


def list_docx_templates() -> list[tuple[str, Path]]:
    """Return ``(display_name, path)`` sorted by name; user ``.iterthink/templates`` overrides bundled stem."""
    bundled: dict[str, Path] = {}
    root = bundled_templates_dir()
    if root.is_dir():
        for p in sorted(root.glob("*.docx")):
            if p.is_file():
                bundled[p.stem.lower()] = p
    user: dict[str, Path] = {}
    udir = user_templates_dir()
    for p in sorted(udir.glob("*.docx")):
        if p.is_file():
            user[p.stem.lower()] = p
    merged = {**bundled, **user}
    paths = sorted(merged.values(), key=lambda x: x.name.lower())
    return [(p.stem, p) for p in paths]


@dataclass
class ExportMeta:
    title_stem: str
    author: str
    date_iso: str
    comment_author: str = ""


@dataclass(frozen=True)
class _ExportStyles:
    """Resolved paragraph/character style names present in the template ``Document``."""

    body: str
    code_block: str
    strong: str | None
    emphasis: str | None
    intense_emphasis: str | None
    code_char: str | None

    def list_paragraph_style(self, doc: Document, *, ordered: bool, depth: int) -> tuple[str, bool]:
        """Return ``(style_name, is_native_list)``. Native list styles use template indentation only."""
        tier = min(max(depth, 0), 2)
        if ordered:
            tier_names = ("List Number", "List Number 2", "List Number 3")
        else:
            tier_names = ("List Bullet", "List Bullet 2", "List Bullet 3")
        for cand in (tier_names[tier], tier_names[0], "List Paragraph", "List"):
            resolved = _style_safe(doc, cand)
            if resolved:
                native = resolved.startswith(
                    ("List Bullet", "List Number", "List Continue", "List Paragraph", "List 2", "List 3")
                ) or resolved == "List"
                return resolved, native
        return self.body, False


def _first_paragraph_style(doc: Document, candidates: tuple[str, ...]) -> str:
    for name in candidates:
        resolved = _style_safe(doc, name)
        if resolved:
            return resolved
    return _style_safe(doc, "Normal") or "Normal"


def _first_character_style(doc: Document, candidates: tuple[str, ...]) -> str | None:
    for name in candidates:
        resolved = _style_resolve(doc, name)
        if not resolved or resolved not in doc.styles:
            continue
        st = doc.styles[resolved]
        if st.type == WD_STYLE_TYPE.CHARACTER:
            return str(st.name)
    return None


def _build_export_styles(doc: Document) -> _ExportStyles:
    body = _first_paragraph_style(doc, ("Body Text", "Text Body", "Normal"))
    code_block = _first_paragraph_style(
        doc,
        ("Preformatted Text", "HTML Preformatted", "No Spacing", "Quote", body, "Normal"),
    )
    return _ExportStyles(
        body=body,
        code_block=code_block,
        strong=_first_character_style(doc, ("Strong",)),
        emphasis=_first_character_style(doc, ("Emphasis",)),
        intense_emphasis=_first_character_style(doc, ("Intense Emphasis",)),
        code_char=_first_character_style(doc, ("Macro Text Char", "HTML Code Char", "Verbatim Char")),
    )


def _style_resolve(doc: Document, name: str) -> str | None:
    """Map UI / casing variants to this document's canonical style name without style_id lookup."""
    tried: list[str] = []
    for k in (name, BabelFish.ui2internal(name), name.lower()):
        if not k or k in tried:
            continue
        tried.append(k)
        if k in doc.styles:
            return str(doc.styles[k].name)
    return None


def _style_safe(doc: Document, name: str) -> str | None:
    """Return canonical **paragraph** style name for ``name``, or None."""
    resolved = _style_resolve(doc, name)
    if not resolved or resolved not in doc.styles:
        return None
    st = doc.styles[resolved]
    if st.type != WD_STYLE_TYPE.PARAGRAPH:
        return None
    return resolved


def _style_list_numpr(doc: Document, style_name: str) -> tuple[int, int] | None:
    """Read ``(num_id, base_ilvl)`` from a paragraph style's ``w:pPr/w:numPr``, or ``None``."""
    if style_name not in doc.styles:
        return None
    st = doc.styles[style_name]
    if st.type != WD_STYLE_TYPE.PARAGRAPH:
        return None
    p_pr = st.element.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    raw_id = num_id_el.get(qn("w:val"))
    if raw_id is None:
        return None
    try:
        num_id = int(raw_id)
    except ValueError:
        return None
    base_ilvl = 0
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if ilvl_el is not None:
        raw_lv = ilvl_el.get(qn("w:val"))
        if raw_lv is not None:
            try:
                base_ilvl = int(raw_lv)
            except ValueError:
                base_ilvl = 0
    return num_id, base_ilvl


def _paragraph_set_list_num_level(paragraph: Any, *, num_id: int, ilvl: int) -> None:
    """Replace paragraph ``w:numPr`` so Word uses ``ilvl`` (0–8) on the given list instance."""
    p_el = paragraph._element
    p_pr = p_el.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:numPr"):
            p_pr.remove(child)
    ilvl_capped = min(max(ilvl, 0), 8)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl_capped))
    num_pr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _heading_style_name(doc: Document, level: int) -> str | None:
    candidates: list[str] = [f"Heading {level}"]
    if level == 1:
        candidates.extend(["Title", "Titel"])
    for cand in candidates:
        s = _style_safe(doc, cand)
        if s:
            return s
    return _first_paragraph_style(doc, ("Normal",))


def _add_paragraph_with_style(doc: Document, style_name: str | None) -> Any:
    if style_name:
        return doc.add_paragraph(style=style_name)
    return doc.add_paragraph()


def _plaintext_from_inline_tokens(tokens: list[Any]) -> str:
    parts: list[str] = []

    def walk(tl: list[Any]) -> None:
        for t in tl:
            if t.type == "text":
                parts.append(t.content or "")
            elif t.type == "softbreak":
                parts.append("\n")
            elif t.type == "footnote_anchor":
                continue
            elif t.children:
                walk(t.children)

    walk(tokens)
    return "".join(parts)


def _list_item_body(tokens: list[Any], i: int) -> tuple[list[Any], int]:
    """Return ``(tokens inside one list item, index after its list_item_close)``.

    Nested ``list_item_open``/``close`` pairs are included in the body so the slice
    does not end at the first inner ``list_item_close``.
    """
    n = len(tokens)
    if i >= n or tokens[i].type != "list_item_open":
        return [], i
    i += 1
    body: list[Any] = []
    depth = 1
    while i < n and depth:
        t = tokens[i]
        if t.type == "list_item_open":
            depth += 1
            body.append(t)
            i += 1
        elif t.type == "list_item_close":
            depth -= 1
            if depth == 0:
                i += 1
                break
            body.append(t)
            i += 1
        else:
            body.append(t)
            i += 1
    return body, i


def _footnote_body_plaintext(block_tokens: list[Any]) -> str:
    parts: list[str] = []

    def walk_block(tl: list[Any]) -> None:
        i = 0
        while i < len(tl):
            t = tl[i]
            if t.type == "paragraph_open":
                i += 1
                if i < len(tl) and tl[i].type == "inline":
                    parts.append(_plaintext_from_inline_tokens(tl[i].children or []))
                    i += 1
                if i < len(tl) and tl[i].type == "paragraph_close":
                    i += 1
                parts.append("\n")
            elif t.type == "heading_open":
                i += 1
                while i < len(tl) and tl[i].type != "heading_close":
                    if tl[i].type == "inline":
                        parts.append(_plaintext_from_inline_tokens(tl[i].children or []))
                    i += 1
                if i < len(tl):
                    i += 1
                parts.append("\n")
            elif t.type in ("bullet_list_open", "ordered_list_open"):
                close = "bullet_list_close" if t.type == "bullet_list_open" else "ordered_list_close"
                i += 1
                while i < len(tl) and tl[i].type != close:
                    walk_block([tl[i]])
                    i += 1
                if i < len(tl):
                    i += 1
            elif t.type == "list_item_open":
                chunk, i = _list_item_body(tl, i)
                walk_block(chunk)
            elif t.type == "fence":
                parts.append((t.content or "").strip() + "\n")
                i += 1
            else:
                i += 1

    walk_block(block_tokens)
    return "".join(parts).strip()


def _split_footnote_block(tokens: list[Any]) -> tuple[list[Any], list[Any]]:
    if not tokens:
        return tokens, []
    if tokens[-1].type != "footnote_block_close":
        return tokens, []
    depth = 0
    start: int | None = None
    for j in range(len(tokens) - 1, -1, -1):
        t = tokens[j]
        if t.type == "footnote_block_close":
            depth += 1
        elif t.type == "footnote_block_open":
            depth -= 1
            if depth == 0:
                start = j
                break
    if start is None:
        return tokens, []
    return tokens[:start], tokens[start:]


def _collect_footnote_defs(fn_block: list[Any]) -> dict[int, list[Any]]:
    out: dict[int, list[Any]] = {}
    if not fn_block or fn_block[0].type != "footnote_block_open":
        return out
    i = 1
    while i < len(fn_block) - 1:
        t = fn_block[i]
        if t.type == "footnote_open":
            pid = int(t.meta["id"])
            i += 1
            chunk: list[Any] = []
            while i < len(fn_block) and fn_block[i].type != "footnote_close":
                chunk.append(fn_block[i])
                i += 1
            out[pid] = chunk
            if i < len(fn_block) and fn_block[i].type == "footnote_close":
                i += 1
        else:
            i += 1
    return out


class _ListRenderMeta(NamedTuple):
    """Markdown list context while walking tokens (``ordered``, nesting ``depth``, 1-based ``item_index``)."""

    ordered: bool
    depth: int
    item_index: int = 0


class _Ctx:
    def __init__(
        self,
        doc: Document,
        md_path: Path,
        fn_id_map: dict[int, int],
        styles: _ExportStyles,
        *,
        paragraph_comments: dict[int, str] | None = None,
        comment_author: str = "",
    ):
        self.doc = doc
        self.md_path = md_path
        self.fn_id_map = fn_id_map
        self.styles = styles
        self.paragraph_comments = paragraph_comments
        self.comment_author = (comment_author or "").strip()
        self.para_comment_idx = 0
        # First block in a list item shows bullet/number when the template has no native list styles.
        self.list_item_mark_pending = True


def _attach_export_paragraph_comment(ctx: _Ctx, paragraph: Any) -> None:
    """One index per top-level block (heading, body paragraph, fence)."""
    if not ctx.paragraph_comments:
        return
    idx = ctx.para_comment_idx
    ctx.para_comment_idx += 1
    text = ctx.paragraph_comments.get(idx)
    if not text or not str(text).strip():
        return
    auth = ctx.comment_author or "iterthink"
    initials = (auth[:2] if len(auth) >= 2 else "it").upper()
    try:
        ctx.doc.add_comment(
            runs=paragraph.runs,
            text=str(text).strip(),
            author=auth[:64],
            initials=initials,
        )
    except Exception:
        pass


def _add_inline_image(ctx: _Ctx, paragraph: Any, t: Any) -> None:
    src = (t.attrs or {}).get("src", "")
    alt = (t.attrs or {}).get("alt", "") or (t.content or "") or "Image"
    path = Path(src)
    if not path.is_absolute():
        path = (ctx.md_path.parent / path).resolve()
    if path.is_file():
        try:
            paragraph.add_run().add_picture(str(path), width=Cm(14))
        except Exception:
            paragraph.add_run(f"[image: {alt}]")
    else:
        paragraph.add_run(f"[missing image: {src}]")


def _render_inline_styled(ctx: _Ctx, paragraph: Any, children: list[Any]) -> None:
    bold = 0
    italic = 0
    i = 0
    while i < len(children):
        t = children[i]
        if t.type == "strong_open":
            bold += 1
            i += 1
        elif t.type == "strong_close":
            bold = max(0, bold - 1)
            i += 1
        elif t.type == "em_open":
            italic += 1
            i += 1
        elif t.type == "em_close":
            italic = max(0, italic - 1)
            i += 1
        elif t.type == "text":
            r = paragraph.add_run(t.content or "")
            if bold and italic:
                if ctx.styles.intense_emphasis:
                    r.style = ctx.styles.intense_emphasis
                else:
                    r.bold = True
                    r.italic = True
            elif bold:
                if ctx.styles.strong:
                    r.style = ctx.styles.strong
                else:
                    r.bold = True
            elif italic:
                if ctx.styles.emphasis:
                    r.style = ctx.styles.emphasis
                else:
                    r.italic = True
            i += 1
        elif t.type == "softbreak":
            # Word ignores newline characters inside <w:t>; use an explicit line break.
            paragraph.add_run().add_break()
            i += 1
        elif t.type == "hardbreak":
            paragraph.add_run().add_break()
            i += 1
        elif t.type == "code_inline":
            r = paragraph.add_run(t.content or "")
            if ctx.styles.code_char:
                r.style = ctx.styles.code_char
            else:
                r.font.name = "Courier New"
                r.font.size = Pt(10)
            i += 1
        elif t.type == "image":
            _add_inline_image(ctx, paragraph, t)
            i += 1
        elif t.type == "footnote_ref":
            pid = int(t.meta["id"])
            wid = ctx.fn_id_map.get(pid)
            if wid is not None:
                docx_footnotes.add_footnote_reference(paragraph, wid)
            i += 1
        elif t.type == "footnote_anchor":
            i += 1
        elif t.type == "link_open":
            href = (t.attrs or {}).get("href", "")
            inner: list[Any] = []
            i += 1
            while i < len(children) and children[i].type != "link_close":
                inner.append(children[i])
                i += 1
            if i < len(children) and children[i].type == "link_close":
                i += 1
            label = _plaintext_from_inline_tokens(inner)
            text = f"{label} ({href})" if href else label
            r = paragraph.add_run(text)
            r.bold = bool(bold)
            r.italic = bool(italic)
            r.font.underline = True
        elif t.type == "inline":
            _render_inline_styled(ctx, paragraph, t.children or [])
            i += 1
        else:
            i += 1


def _list_depth_parent(list_meta: _ListRenderMeta | None) -> int:
    return list_meta.depth if list_meta is not None else -1


def _render_blocks(ctx: _Ctx, tokens: list[Any], list_meta: _ListRenderMeta | None = None) -> None:
    i = 0
    n = len(tokens)
    while i < n:
        t = tokens[i]
        if t.type == "footnote_block_open":
            depth = 1
            i += 1
            while i < n and depth:
                if tokens[i].type == "footnote_block_open":
                    depth += 1
                elif tokens[i].type == "footnote_block_close":
                    depth -= 1
                i += 1
            continue
        if t.type == "heading_open":
            level = int(t.tag[1]) if t.tag and t.tag.startswith("h") else 1
            level = max(1, min(level, 9))
            i += 1
            style = _heading_style_name(ctx.doc, level)
            p = _add_paragraph_with_style(ctx.doc, style)
            if i < n and tokens[i].type == "inline":
                _render_inline_styled(ctx, p, tokens[i].children or [])
                i += 1
            if i < n and tokens[i].type == "heading_close":
                i += 1
            _attach_export_paragraph_comment(ctx, p)
            continue
        if t.type == "paragraph_open":
            i += 1
            if list_meta is not None:
                ordered, depth, item_index = list_meta
                st_name, native = ctx.styles.list_paragraph_style(ctx.doc, ordered=ordered, depth=depth)
                p = ctx.doc.add_paragraph(style=st_name)
                applied_style_num = False
                if native:
                    snum = _style_list_numpr(ctx.doc, st_name)
                    if snum is not None:
                        num_id, _base_ilvl = snum
                        _paragraph_set_list_num_level(p, num_id=num_id, ilvl=depth)
                        applied_style_num = True
                if not applied_style_num:
                    p.paragraph_format.left_indent = Cm(0.55 * (depth + 1))
                    p.paragraph_format.first_line_indent = Cm(-0.32)
                    if ctx.list_item_mark_pending:
                        if ordered and item_index > 0:
                            p.add_run(f"{item_index}.\t")
                        elif not ordered:
                            p.add_run("•\t")
                if ctx.list_item_mark_pending:
                    ctx.list_item_mark_pending = False
            else:
                p = ctx.doc.add_paragraph(style=ctx.styles.body)
            if i < n and tokens[i].type == "inline":
                _render_inline_styled(ctx, p, tokens[i].children or [])
                i += 1
            if i < n and tokens[i].type == "paragraph_close":
                i += 1
            if list_meta is None:
                _attach_export_paragraph_comment(ctx, p)
            continue
        if t.type == "fence":
            body = (t.content or "").rstrip("\n")
            p = ctx.doc.add_paragraph(style=ctx.styles.code_block)
            if list_meta is not None:
                ordered, depth, _idx = list_meta
                _st, native = ctx.styles.list_paragraph_style(ctx.doc, ordered=ordered, depth=depth)
                if not native:
                    p.paragraph_format.left_indent = Cm(0.55 * (depth + 1))
                    p.paragraph_format.first_line_indent = Cm(-0.32)
                ctx.list_item_mark_pending = False
            lines = body.split("\n")
            for li, line in enumerate(lines):
                if li:
                    p.add_run().add_break()
                r = p.add_run(line)
                if ctx.styles.code_char:
                    r.style = ctx.styles.code_char
                else:
                    r.font.name = "Courier New"
                    r.font.size = Pt(10)
            _attach_export_paragraph_comment(ctx, p)
            i += 1
            continue
        if t.type == "bullet_list_open":
            parent = _list_depth_parent(list_meta)
            i += 1
            while i < n and tokens[i].type != "bullet_list_close":
                if tokens[i].type == "list_item_open":
                    inner, i = _list_item_body(tokens, i)
                    ctx.list_item_mark_pending = True
                    _render_blocks(ctx, inner, _ListRenderMeta(False, parent + 1, 0))
                else:
                    i += 1
            if i < n and tokens[i].type == "bullet_list_close":
                i += 1
            continue
        if t.type == "ordered_list_open":
            parent = _list_depth_parent(list_meta)
            i += 1
            ord_n = 0
            while i < n and tokens[i].type != "ordered_list_close":
                if tokens[i].type == "list_item_open":
                    ord_n += 1
                    inner, i = _list_item_body(tokens, i)
                    ctx.list_item_mark_pending = True
                    _render_blocks(ctx, inner, _ListRenderMeta(True, parent + 1, ord_n))
                else:
                    i += 1
            if i < n and tokens[i].type == "ordered_list_close":
                i += 1
            continue
        if t.type == "hr":
            i += 1
            ctx.doc.add_paragraph("— — —", style=ctx.styles.body)
            continue
        if t.type == "html_block":
            i += 1
            continue
        i += 1


def markdown_to_docx(
    *,
    markdown_src: str,
    md_path: Path,
    template_path: Path,
    output_path: Path,
    meta: ExportMeta,
    paragraph_comments: dict[int, str] | None = None,
) -> None:
    """Write ``output_path`` from ``markdown_src`` using ``template_path``."""
    from markdown_it import MarkdownIt
    from mdit_py_plugins.footnote import footnote_plugin

    # breaks=True: single newlines in the markdown file become softbreak tokens and render as
    # Word line breaks (CommonMark otherwise treats them like spaces in one paragraph).
    md = MarkdownIt("commonmark", {"breaks": True}).use(footnote_plugin)
    tokens = md.parse(markdown_src)
    main_toks, fn_block = _split_footnote_block(tokens)
    defs = _collect_footnote_defs(fn_block)

    doc = Document(str(template_path))
    fn_part = docx_footnotes.ensure_footnotes_part(doc)
    next_wid = max(docx_footnotes.max_footnote_id(fn_part), 0) + 1
    fn_map: dict[int, int] = {}
    for pid in sorted(defs.keys()):
        fn_map[pid] = next_wid
        next_wid += 1

    for pid in sorted(defs.keys()):
        plain = _footnote_body_plaintext(defs[pid])
        blob = docx_footnotes.paragraph_xml_plain(plain)
        docx_footnotes.append_footnote_xml(fn_part, fn_map[pid], [blob])

    ctx = _Ctx(
        doc,
        md_path.resolve(),
        fn_map,
        _build_export_styles(doc),
        paragraph_comments=paragraph_comments,
        comment_author=meta.comment_author or meta.author,
    )
    _render_blocks(ctx, main_toks, None)

    mapping = {
        "{Titel}": meta.title_stem,
        "{Date}": meta.date_iso,
        "{Author}": meta.author,
        # Same value as export "Author" in settings (templates often use {Name}).
        "{Name}": meta.author,
    }
    docx_placeholders.apply_docx_placeholders(doc, mapping)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
