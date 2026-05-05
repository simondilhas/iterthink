"""Convert DOCX / PDF imports to Markdown with assets beside the target document."""

from __future__ import annotations

import hashlib
import logging
import os
import re
import statistics
from pathlib import Path
from typing import Literal

from iterthink import config

logger = logging.getLogger(__name__)

PdfProfileHeuristic = Literal["text", "plan"]

ALLOWED_IMPORT_EXTENSIONS = frozenset({"docx", "pdf"})


def validate_extension(path: Path) -> str | None:
    """Return lower-case extension if allowed, else None."""
    suf = path.suffix.lower().lstrip(".")
    return suf if suf in ALLOWED_IMPORT_EXTENSIONS else None


def docx_to_markdown(src: Path, asset_dir: Path) -> str:
    """Extract markdown from Word; images written under ``asset_dir`` with relative links."""
    from docx import Document as DocxDocument

    asset_dir.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument(str(src))
    markdown_lines: list[str] = []
    stem_safe = re.sub(r"[^\w\-]+", "_", src.stem)[:80] or "doc"
    image_counter = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        name = para.style.name if para.style else ""
        if not text:
            markdown_lines.append("")
            continue
        if name.startswith("Heading 1"):
            markdown_lines.append(f"# {text}")
        elif name.startswith("Heading 2"):
            markdown_lines.append(f"## {text}")
        elif name.startswith("Heading 3"):
            markdown_lines.append(f"### {text}")
        else:
            markdown_lines.append(text)

    for shape in doc.inline_shapes:
        try:
            image_part = shape._inline.graphic.graphicData.pic.blip.fill.blip
            r_id = image_part.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_id:
                continue
            related_part = doc.part.related_parts[r_id]
            image_bytes = related_part.blob
            content_type = getattr(related_part, "content_type", "") or "image/png"
            ext = content_type.split("/")[-1].lower()
            if ext in ("jpeg", "jpg"):
                ext = "jpg"
            elif ext != "png":
                ext = "png"
            image_filename = f"{stem_safe}_image_{image_counter}.{ext}"
            image_path = asset_dir / image_filename
            image_path.write_bytes(image_bytes)
            rel = os.path.relpath(image_path, asset_dir.parent)
            markdown_lines.append(f"\n![Image]({rel})\n")
            image_counter += 1
        except Exception as ex:
            logger.debug("Inline image skipped: %s", ex)

    parts: list[str] = []
    buf: list[str] = []
    for line in markdown_lines:
        if line == "":
            if buf:
                parts.append("\n".join(buf))
                buf = []
            continue
        buf.append(line)
    if buf:
        parts.append("\n".join(buf))

    return "\n\n".join(parts) if parts else ""


# Lines that are only a bullet glyph (text often continues on the next line).
_BULLET_ONLY_LINE = re.compile(r"^[•·▪▸►◦‣⁃]\s*$")
# Inline bullet + content on same line
_INLINE_BULLET = re.compile(r"^[•·▪▸►◦]\s*(.+)$")
_MARKDOWNISH_BULLET = re.compile(r"^[\-\*\+]\s+(.+)$")
_NUMBERED_LINE = re.compile(r"^(\d{1,3})[\.\)]\s+(.*)$")


def _strip_inline_bullet(text: str) -> str | None:
    m = _INLINE_BULLET.match(text)
    if m:
        return m.group(1).strip()
    m = _MARKDOWNISH_BULLET.match(text)
    if m:
        return m.group(1).strip()
    return None


def _pdf_dict_to_markdown(src: Path) -> str:
    """
    Extract PDF using ``Page.get_text(\"dict\")``: reading order, font-size headings,
    bullets / numbered lists, and paragraph merging from vertical gaps.
    """
    import fitz

    doc = fitz.open(str(src))
    page_chunks: list[str] = []
    try:
        all_sizes: list[float] = []
        for pi in range(len(doc)):
            td = doc[pi].get_text("dict") or {}
            for b in td.get("blocks", []):
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for sp in line.get("spans", []):
                        z = sp.get("size") or 0
                        if z and float(z) > 0:
                            all_sizes.append(float(z))
        if len(all_sizes) >= 6:
            sorted_sz = sorted(all_sizes)
            body_med = sorted_sz[len(sorted_sz) // 4]
        else:
            body_med = statistics.median(all_sizes) if all_sizes else 11.0
        body_med = max(6.0, min(body_med, 22.0))

        for pi in range(len(doc)):
            page_num = pi + 1
            page = doc[pi]
            td = page.get_text("dict") or {}
            blocks = [b for b in td.get("blocks", []) if b.get("type") == 0]
            flat: list[tuple[float, float, float, float, dict]] = []
            for b in blocks:
                for line in b.get("lines", []):
                    bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
                    x0, y0, _x1, y1 = bbox[0], bbox[1], bbox[2], bbox[3]
                    flat.append((y0, x0, y1, line))
            flat.sort(key=lambda t: (round(t[0], 2), round(t[1], 2)))

            events: list[tuple[str, tuple]] = []
            pending_bullet = False

            for _y0s, _x0s, _y1s, line in flat:
                spans = line.get("spans") or []
                if not spans:
                    continue
                raw = "".join(sp.get("text", "") for sp in spans)
                text = raw.strip()
                if not text:
                    continue

                if _BULLET_ONLY_LINE.match(text):
                    pending_bullet = True
                    continue

                bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
                ly0, ly1 = float(bbox[1]), float(bbox[3])

                if pending_bullet:
                    events.append(("bullet", (text, ly0, ly1)))
                    pending_bullet = False
                    continue

                nm = _NUMBERED_LINE.match(text)
                if nm:
                    events.append(("ol", (nm.group(1), nm.group(2).strip(), ly0, ly1)))
                    continue

                stripped = _strip_inline_bullet(text)
                if stripped is not None:
                    events.append(("bullet", (stripped, ly0, ly1)))
                    continue

                sizes = [float(sp.get("size") or 0) for sp in spans if (sp.get("size") or 0) > 0]
                max_sz = max(sizes) if sizes else body_med
                ratio = max_sz / body_med if body_med > 0 else 1.0

                long_line = len(text) > 180
                strong_heading = ratio >= 1.55

                if not long_line or strong_heading:
                    if ratio >= 1.48:
                        events.append(("h1", (text, ly0, ly1)))
                        continue
                    if ratio >= 1.26:
                        events.append(("h2", (text, ly0, ly1)))
                        continue
                    if ratio >= 1.11:
                        events.append(("h3", (text, ly0, ly1)))
                        continue

                events.append(("body", (text, ly0, ly1)))

            md_body = _pdf_events_to_markdown(events, body_med)
            marker = f"<!-- page:{page_num} -->"
            if md_body.strip():
                page_chunks.append(f"{marker}\n\n{md_body.strip()}")
            else:
                page_chunks.append(marker)
    finally:
        doc.close()

    return "\n\n".join(page_chunks)


def _pdf_events_to_markdown(events: list[tuple[str, tuple]], body_med: float) -> str:
    """Turn classified line events into Markdown (paragraphs, lists, headings)."""
    gap_para = max(4.0, body_med * 0.55)
    gap_join = max(2.5, body_med * 0.38)

    parts: list[str] = []
    cur_para: list[str] = []
    prev_y1: float | None = None
    list_lines: list[str] = []

    def flush_para() -> None:
        nonlocal cur_para, prev_y1
        if cur_para:
            parts.append("\n".join(cur_para))
        cur_para = []
        prev_y1 = None

    def flush_list() -> None:
        nonlocal list_lines
        if list_lines:
            parts.append("\n".join(list_lines))
        list_lines = []

    for ev in events:
        kind, payload = ev
        if kind == "body":
            flush_list()
            text = str(payload[0])
            y0, y1 = float(payload[1]), float(payload[2])
            if prev_y1 is not None and cur_para and (y0 - prev_y1) <= gap_join:
                cur_para[-1] = (cur_para[-1].rstrip() + " " + text.lstrip()).strip()
            elif prev_y1 is not None and cur_para and (y0 - prev_y1) > gap_para:
                flush_para()
                cur_para.append(text)
            else:
                if cur_para:
                    parts.append("\n".join(cur_para))
                    cur_para = []
                cur_para.append(text)
            prev_y1 = y1
            continue

        flush_para()
        if kind == "h1":
            flush_list()
            parts.append(f"# {payload[0]}")
        elif kind == "h2":
            flush_list()
            parts.append(f"## {payload[0]}")
        elif kind == "h3":
            flush_list()
            parts.append(f"### {payload[0]}")
        elif kind == "bullet":
            list_lines.append(f"- {str(payload[0])}")
        elif kind == "ol":
            list_lines.append(f"{payload[0]}. {payload[1]}")

    flush_para()
    flush_list()

    return "\n\n".join(parts)


def _pdf_to_markdown_legacy(src: Path) -> str:
    """
    Plain PyMuPDF ``get_text`` extraction (fallback).

    Inserts ``<!-- page:N -->`` (1-based) before each page's content for scroll mapping.
    """
    import fitz

    doc = fitz.open(str(src))
    chunks: list[str] = []
    try:
        for page_index in range(len(doc)):
            page_num = page_index + 1
            page = doc[page_index]
            raw = page.get_text("text") or ""
            paras = [p.strip() for p in raw.split("\n\n") if p.strip()]
            if not paras:
                blocks = page.get_text("blocks") or []
                lines: list[str] = []
                for b in blocks:
                    if isinstance(b, (list, tuple)) and len(b) >= 5:
                        t = b[4]
                        if isinstance(t, str) and t.strip():
                            lines.append(t.strip())
                raw2 = "\n".join(lines)
                paras = [p.strip() for p in raw2.split("\n\n") if p.strip()]
                if not paras and lines:
                    paras = [" ".join(lines)]
            body = "\n\n".join(paras) if paras else ""
            marker = f"<!-- page:{page_num} -->"
            if body.strip():
                chunks.append(f"{marker}\n\n{body}")
            else:
                chunks.append(marker)
    finally:
        doc.close()

    return "\n\n".join(chunks)


def pdf_to_markdown(src: Path) -> str:
    """
    PDF → Markdown with ``<!-- page:N -->`` markers.

    Primary path: structured PyMuPDF ``dict`` extraction (reading order, font-size headings,
    bullets / numbered lists, paragraph gaps). Optional: PyMuPDF4LLM. Last resort: plain text.
    """

    def _has_body(md: str) -> bool:
        raw = re.sub(r"<!--\s*page:\d+\s*-->", "", md)
        return bool(raw.strip())

    try:
        md_struct = _pdf_dict_to_markdown(src)
        if _has_body(md_struct):
            return md_struct
    except Exception as ex:
        logger.warning("Structured PDF extraction failed: %s", ex)

    src_abs = str(src.resolve())
    chunks_out: list[str] | None = None

    try:
        import pymupdf4llm  # type: ignore[import-untyped]

        result = pymupdf4llm.to_markdown(
            src_abs,
            page_chunks=True,
            write_images=False,
            embed_images=False,
        )
        if isinstance(result, list):
            chunks_out = []
            for chunk in result:
                meta = chunk.get("metadata") or {}
                page_num = meta.get("page_number")
                if not isinstance(page_num, int):
                    page_num = len(chunks_out) + 1
                body = (chunk.get("text") or "").strip()
                marker = f"<!-- page:{page_num} -->"
                if body:
                    chunks_out.append(f"{marker}\n\n{body}")
                else:
                    chunks_out.append(marker)
    except ImportError:
        logger.debug("pymupdf4llm not available")
    except Exception as ex:
        logger.warning("PyMuPDF4LLM PDF extraction failed: %s", ex)

    if chunks_out:
        md = "\n\n".join(chunks_out)
        if _has_body(md):
            return md

    md = _pdf_to_markdown_legacy(src)
    if not md.strip():
        return (
            "<!-- PDF text extraction returned no content. This may be a scanned PDF. -->\n\n"
            "*Note: No extractable text; use Original PDF in Compare to view.*"
        )
    return md


def classify_pdf_profile(src: Path) -> PdfProfileHeuristic:
    """
    Rough split: text-heavy PDFs → ``text``; sparse extraction → ``plan`` / drawing-like.
    """
    import fitz

    doc = fitz.open(str(src))
    try:
        n = len(doc)
        if n == 0:
            return "plan"
        total_chars = 0
        for i in range(n):
            total_chars += len((doc[i].get_text() or "").strip())
        avg = total_chars / max(n, 1)
        # Few readable characters per page → treat as drawing-first.
        if avg < 120:
            return "plan"
        return "text"
    finally:
        doc.close()


def pdf_render_cache_dir(pdf_abs: Path) -> Path:
    """Stable cache folder under store for rendered PNG pages."""
    st = pdf_abs.stat()
    key_src = f"{pdf_abs.resolve()}:{st.st_mtime_ns}:{st.st_size}".encode()
    key = hashlib.sha256(key_src).hexdigest()[:24]
    d = config.STORE_DIR / "pdf_render_cache" / key
    d.mkdir(parents=True, exist_ok=True)
    return d


def render_pdf_to_png_pages(pdf_abs: Path) -> list[Path]:
    """
    Rasterize each PDF page to PNG under the render cache.
    Returns ordered list of PNG paths.
    """
    import fitz

    cache = pdf_render_cache_dir(pdf_abs)
    marker = cache / ".source"
    src_tag = f"{pdf_abs.resolve()}:{pdf_abs.stat().st_mtime_ns}"
    if marker.is_file() and marker.read_text(encoding="utf-8") == src_tag:
        existing = sorted(cache.glob("page_*.png"))
        if existing:
            return existing

    for old in cache.glob("page_*.png"):
        old.unlink(missing_ok=True)

    doc = fitz.open(str(pdf_abs))
    out: list[Path] = []
    try:
        for i in range(len(doc)):
            page = doc[i]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            p = cache / f"page_{i + 1:04d}.png"
            pix.save(str(p))
            out.append(p)
    finally:
        doc.close()

    marker.write_text(src_tag, encoding="utf-8")
    return out


def plain_text_preview_from_docx(path: Path, *, max_chars: int = 600_000) -> str:
    """Paragraph text from a Word file for Compare preview (plain text, not Markdown)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    body = "\n\n".join(parts)
    if len(body) > max_chars:
        body = body[:max_chars].rstrip() + "\n\n…"
    return body if body.strip() else "(No paragraph text in document)"


def import_file_to_markdown(src: Path, md_path: Path) -> str:
    """Route by extension. DOCX assets go to ``<md_stem>_assets/`` next to ``md_path``."""
    ext = validate_extension(src)
    if ext is None:
        raise ValueError(f"Unsupported import type: {src.suffix}")

    if ext == "docx":
        asset_dir = md_path.parent / f"{md_path.stem}_assets"
        return docx_to_markdown(src, asset_dir)

    return pdf_to_markdown(src)

