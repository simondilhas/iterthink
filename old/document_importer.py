"""Document import functionality for Word documents and PDFs."""
from docx import Document as DocxDocument
from pathlib import Path
import os
import re
from PIL import Image
import io
import logging
import shutil
from typing import Optional, Sequence, Union

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


ExtractedImage = Union[str, tuple[str, str]]


def store_pdf_assets_locally(
    pdf_source_path: str,
    doc_id: int,
    version_id: int,
    extracted_images: Optional[Sequence[ExtractedImage]] = None,
) -> Optional[str]:
    """
    Persist the original PDF (and any extracted images) to the local data directory
    for environments without cloud storage.
    """

    try:
        scan_dir = Path("data/_scans") / str(doc_id) / str(version_id)
        scan_dir.mkdir(parents=True, exist_ok=True)
        pdf_destination = scan_dir / "original.pdf"
        shutil.copy2(pdf_source_path, pdf_destination)
        logger.info("Stored original PDF locally at %s", pdf_destination)

        if extracted_images:
            images_dir = scan_dir / "embedded_images"
            images_dir.mkdir(parents=True, exist_ok=True)
            for idx, img_data in enumerate(extracted_images, start=1):
                if isinstance(img_data, tuple):
                    src_path, clean_name = img_data
                    filename = clean_name or Path(src_path).name or f"image_{idx:04d}.png"
                else:
                    src_path = img_data
                    filename = Path(src_path).name or f"image_{idx:04d}.png"

                destination = images_dir / filename
                try:
                    shutil.copy2(src_path, destination)
                    logger.debug("Stored embedded image locally at %s", destination)
                except Exception as img_exc:
                    logger.warning("Failed to store embedded image '%s': %s", filename, img_exc)

        return str(pdf_destination)
    except Exception as exc:
        logger.error("Failed to store PDF assets locally: %s", exc, exc_info=True)
        return None


def _cleanup_extracted_images(extracted_images: Sequence[ExtractedImage]) -> None:
    for img_data in extracted_images or []:
        try:
            persistent_path = img_data[0] if isinstance(img_data, tuple) else img_data
            Path(persistent_path).unlink(missing_ok=True)
        except Exception as cleanup_exc:
            logger.debug("Failed to cleanup temporary image file: %s", cleanup_exc)


def import_document(file_path, file_ext, title, db, user_id=None, project=None, tag=None, content_type=None, status=None):
    """Import a document (Word or PDF) and convert to markdown with images.
    
    Args:
        file_path: Path to the document file
        file_ext: File extension (pdf, doc, docx)
        title: Document title
        db: Database session
        user_id: User ID (optional)
        project: Project name (optional)
        tag: Tag name (optional)
        content_type: Content type (law, standard, contract, etc.) - optional, can be classified automatically
        status: Document status (draft, shared, published, archived) - optional, defaults to "draft"
    """
    
    # Import services
    from app_server.services.openstack_storage import openstack_storage
    from app_server.services.converters.pdf_converter import pdf_converter
    
    # Convert to markdown based on file type
    if file_ext in ['doc', 'docx']:
        markdown_content = convert_word_to_markdown(file_path)
        doc_type = "docx"
        total_pages = None  # DOCX doesn't have explicit pages
    elif file_ext == 'pdf':
        logger.info(f"Starting PDF conversion for {file_path}")
        extracted_images = []  # Initialize
        try:
            markdown_content, page_breaks, paragraph_page_map, extracted_images = convert_pdf_to_markdown(file_path)
            logger.info(f"PDF conversion completed: {len(markdown_content)} chars extracted, {len(paragraph_page_map)} paragraphs mapped to pages, {len(extracted_images)} embedded images extracted")
            
            # Check if we got any text
            if not markdown_content or not markdown_content.strip():
                logger.warning(f"⚠️  PDF text extraction returned empty content. This PDF may be image-based (scanned) and require OCR.")
                # Still create document but with empty content placeholder
                markdown_content = "<!-- PDF text extraction returned no content. This may be a scanned/image-based PDF. -->\n\n*Note: This PDF could not be automatically converted to text. The original PDF is available for viewing.*"
                paragraph_page_map = []
        except Exception as e:
            logger.error(f"Error during PDF conversion: {e}", exc_info=True)
            # Create placeholder content
            markdown_content = f"<!-- Error extracting text from PDF: {str(e)} -->\n\n*Note: Text extraction failed. The original PDF is available for viewing.*"
            paragraph_page_map = []
            page_breaks = []
            extracted_images = []
        
        doc_type = "pdf"
        total_pages = pdf_converter.get_page_count(file_path)
        logger.info(f"PDF has {total_pages} pages")

        if extracted_images:
            logger.info(f"Processing {len(extracted_images)} extracted images for markdown references")
            image_ref_pattern = re.compile(r'!\[([^\]]*)\]\((image_\d{4}\.png)\)')
            image_refs = image_ref_pattern.findall(markdown_content)
            logger.info(f"Found {len(image_refs)} image references in markdown")

            if not image_refs:
                logger.warning("No image references matched the expected pattern; extracted images may not be linked.")
                permissive_pattern = re.compile(r'!\[([^\]]*)\]\((.*?\.png)\)')
                permissive_refs = permissive_pattern.findall(markdown_content)
                if permissive_refs:
                    logger.info(f"Found {len(permissive_refs)} image references with permissive pattern: {permissive_refs[:5]}")
            else:
                replacement_count = [0]

                def replace_image_with_relative_path(match):
                    alt_text = match.group(1)
                    image_filename = match.group(2)
                    relative_path = f"embedded_images/{image_filename}"
                    replacement_count[0] += 1
                    return f"![{alt_text}]({relative_path})"

                old_markdown = markdown_content
                markdown_content = image_ref_pattern.sub(replace_image_with_relative_path, markdown_content)
                logger.info("Replaced %s image references with relative embedded paths", replacement_count[0])

                if old_markdown == markdown_content:
                    logger.error("Markdown content was not updated when replacing embedded image references")
    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Only Word documents (.doc, .docx) and PDFs (.pdf) are supported.")
    
    # Create document
    from app_server.db.models import Document, Version
    # Use provided project or default to "my-ideas"
    project_name = project if project and project.strip() else "my-ideas"
    tag_name = tag if tag and tag.strip() else None
    content_type_value = content_type.strip() if content_type and content_type.strip() else None
    status_value = status.strip() if status and status.strip() else "draft"
    logger.info(f"Creating document with content_type: {repr(content_type_value)}, status: {repr(status_value)}")
    doc = Document(
        title=title,
        project=project_name,
        tag=tag_name,
        document_type=doc_type,
        content_type=content_type_value,
        status=status_value,
        owner_id=user_id
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    logger.info(f"Document created with id={doc.id}, content_type={repr(doc.content_type)}")
    
    # Create initial version
    initial_version = Version(
        document_id=doc.id,
        title="first version",
        commit_message="first version",
        source="user",
        created_by=user_id
    )
    db.add(initial_version)
    db.commit()
    db.refresh(initial_version)
    
    # For PDFs: Upload original file and page images to cloud storage
    if file_ext == 'pdf':
        logger.info("OpenStack storage enabled: %s", openstack_storage.is_enabled())
        pdf_stored = False
        if openstack_storage.is_enabled():
            try:
                # Upload original PDF
                original_blob_name = f"{doc.id}/{initial_version.id}/original.pdf"
                logger.info("Uploading original PDF to %s", original_blob_name)
                result = openstack_storage.upload_file(file_path, original_blob_name)
                logger.info("Upload result: %s", result)
                pdf_stored = bool(result)

                if not pdf_stored:
                    logger.warning("OpenStack upload returned no URL; will store PDF locally as fallback.")
                
                # Upload extracted embedded images from PDF
                if extracted_images:
                    logger.info(f"Uploading {len(extracted_images)} embedded images from PDF...")
                    
                    # Upload extracted embedded images from PDF
                    for img_idx, img_data in enumerate(extracted_images, start=1):
                        try:
                            # Handle tuple (persistent_path, clean_name) or string path (backward compatibility)
                            if isinstance(img_data, tuple):
                                persistent_path, clean_name = img_data
                            else:
                                persistent_path = img_data
                                clean_name = None
                            
                            with open(persistent_path, 'rb') as img_file:
                                img_bytes = img_file.read()
                            img_blob_name = f"{doc.id}/{initial_version.id}/embedded_images/image_{img_idx:04d}.png"
                            logger.info("Uploading embedded image %s to %s", img_idx, img_blob_name)
                            openstack_storage.upload_bytes(img_bytes, img_blob_name, "image/png")
                        except Exception as e:
                            logger.warning(f"Failed to upload embedded image {img_idx}: {e}")
                    
                    logger.info("All embedded images uploaded successfully")
            except Exception as e:
                # Log error but don't fail the import
                logger.error("Failed to upload PDF to OpenStack storage: %s", e, exc_info=True)
                pdf_stored = False
        if not pdf_stored:
            logger.info("Storing original PDF locally for doc_id=%s version_id=%s", doc.id, initial_version.id)
            store_pdf_assets_locally(file_path, doc.id, initial_version.id, extracted_images)

        _cleanup_extracted_images(extracted_images)
    
    # Save content to database
    logger.info(f"Saving markdown content to database (length: {len(markdown_content)} chars)")
    # Log a sample to verify it has the expected image references
    content_sample = markdown_content[:500] if len(markdown_content) > 500 else markdown_content
    if "image_" in content_sample and "http" not in content_sample:
        logger.warning("WARNING: Markdown sample contains 'image_' but no 'http' - signed URLs may not be present!")
    logger.info(f"Content sample to be saved (first 500 chars): {content_sample}")
    initial_version.content = markdown_content
    logger.info(f"Content assigned to initial_version. About to commit...")
    
    # Assign paragraph IDs (no parent, so all new)
    from app_server.services.services import assign_paragraph_ids
    from app_server.db.models import VersionParagraph
    
    para_id_map = assign_paragraph_ids(doc.id, None, markdown_content, db)
    
    # Create mapping from paragraph index to PDF page number (only for PDFs)
    para_index_to_page = {}
    if file_ext == 'pdf' and paragraph_page_map:
        # paragraph_page_map is list of (markdown_line_index, page_number)
        # We need to map this to actual paragraph indices after splitting
        from app_server.services.services import split_paragraphs
        paras = split_paragraphs(markdown_content)
        
        # Build a mapping from markdown character position to paragraph index
        # Split markdown by double newlines (paragraph boundaries) to get approximate positions
        markdown_parts = markdown_content.split('\n\n')
        para_start_positions = []
        current_pos = 0
        for part in markdown_parts:
            para_start_positions.append(current_pos)
            current_pos += len(part) + 2  # +2 for the \n\n
        
        # Map paragraph_page_map indices (which are line indices in markdown_lines)
        # to actual paragraph indices
        # We'll use a simple approach: map based on the order of paragraphs in the markdown
        for markdown_line_idx, page_num in paragraph_page_map:
            # Find the paragraph that contains this markdown line
            # Since markdown_line_idx is the index in the markdown_lines list,
            # and we split paragraphs by \n\n, we need to approximate
            para_idx = min(markdown_line_idx, len(paras) - 1) if paras else 0
            if para_idx not in para_index_to_page or para_index_to_page[para_idx] is None:
                para_index_to_page[para_idx] = page_num
    
    for para_index, para_id in para_id_map.items():
        # Get PDF page number for this paragraph if available
        pdf_page = para_index_to_page.get(para_index)
        
        vp = VersionParagraph(
            version_id=initial_version.id,
            paragraph_id=para_id,
            paragraph_index=para_index,
            pdf_page_number=pdf_page
        )
        db.add(vp)
    
    db.commit()
    
    # If PDF: extract and persist page dimensions and external annotations (comments)
    if file_ext == 'pdf':
        try:
            pages = pdf_converter.extract_page_dimensions(file_path)
            annots = pdf_converter.extract_annotations(file_path)
            boxes_by_page = pdf_converter.extract_paragraph_boxes(file_path)
            from app_server.db.models import VersionPage, ExternalAnnotation, VersionParagraph, VersionParagraphBox, Comment, Document as DocModel, Version as VersionModel
            from app_server.services.projects import get_project_encryption_key_from_document
            from app_server.core.encryption_utils import encrypt_project_data
            
            # Persist pages
            for p in pages:
                vp = VersionPage(
                    version_id=initial_version.id,
                    page_number=p["page_number"],
                    width_pt=p["width_pt"],
                    height_pt=p["height_pt"],
                )
                db.add(vp)
            
            # Build paragraph boxes by aligning page text blocks (reading-order) to paragraphs on that page
            vps_on_version = db.query(VersionParagraph).filter(VersionParagraph.version_id == initial_version.id).all()
            vps_by_page = {}
            for vp in sorted(vps_on_version, key=lambda x: (x.pdf_page_number or 10**9, x.paragraph_index)):
                if vp.pdf_page_number is None:
                    continue
                vps_by_page.setdefault(vp.pdf_page_number, []).append(vp)
            # Persist paragraph boxes
            for page_num, vps_list in vps_by_page.items():
                blocks = boxes_by_page.get(page_num, [])
                # assign in order, min length
                assign_count = min(len(vps_list), len(blocks))
                for i in range(assign_count):
                    vp = vps_list[i]
                    blk = blocks[i]
                    vpb = VersionParagraphBox(
                        version_id=initial_version.id,
                        paragraph_id=vp.paragraph_id,
                        page_number=page_num,
                        x0_pt=blk["x0_pt"], y0_pt=blk["y0_pt"],
                        x1_pt=blk["x1_pt"], y1_pt=blk["y1_pt"],
                        x0=blk["x0"], y0=blk["y0"],
                        x1=blk["x1"], y1=blk["y1"],
                        read_order_index=blk.get("order"),
                    )
                    db.add(vpb)
            db.flush()
            
            # Build nearest-paragraph mapper using paragraph boxes
            boxes_index = {}
            for page_num, vps_list in vps_by_page.items():
                boxes_index[page_num] = []
            for pb in db.query(VersionParagraphBox).filter(VersionParagraphBox.version_id == initial_version.id).all():
                boxes_index.setdefault(pb.page_number, []).append(pb)
            
            # Get encryption key for comment messages
            project_key = get_project_encryption_key_from_document(doc.id, db)
            
            # Persist annotations and create linked comments (map to nearest bbox on same page)
            for a in annots:
                page_num = a["page_number"]
                mapped_para_id = None
                nearest_box = None
                candidates = boxes_index.get(page_num, [])
                if candidates:
                    # center of annotation
                    ax = (a["x0_pt"] + a["x1_pt"]) / 2.0
                    ay = (a["y0_pt"] + a["y1_pt"]) / 2.0
                    best_d = None
                    for pb in candidates:
                        cx = (pb.x0_pt + pb.x1_pt) / 2.0
                        cy = (pb.y0_pt + pb.y1_pt) / 2.0
                        d = (ax - cx) * (ax - cx) + (ay - cy) * (ay - cy)
                        if best_d is None or d < best_d:
                            best_d = d
                            nearest_box = pb
                    if nearest_box:
                        mapped_para_id = nearest_box.paragraph_id
                # fallback: pick first paragraph on page if no boxes
                if not mapped_para_id and page_num in vps_by_page and vps_by_page[page_num]:
                    mapped_para_id = vps_by_page[page_num][0].paragraph_id
                
                # Derive a comment message from annotation payload (best-effort)
                import json as _json
                payload_text = ""
                try:
                    payload = _json.loads(a.get("payload_json") or "{}")
                    payload_text = payload.get("content") or payload.get("subject") or payload.get("title") or ""
                except Exception:
                    payload_text = ""
                
                comment_id = None
                if mapped_para_id is not None:
                    encrypted_message = encrypt_project_data(payload_text, project_key if project_key else None) if payload_text else ""
                    comment = Comment(
                        version_id=initial_version.id,
                        paragraph_id=mapped_para_id,
                        user_id=None,
                        tag="external",
                        message=encrypted_message,
                        status="open",
                    )
                    db.add(comment)
                    db.flush()
                    comment_id = comment.id
                
                ea = ExternalAnnotation(
                    version_id=initial_version.id,
                    comment_id=comment_id,
                    page_number=page_num,
                    x0_pt=a["x0_pt"],
                    y0_pt=a["y0_pt"],
                    x1_pt=a["x1_pt"],
                    y1_pt=a["y1_pt"],
                    x0=a["x0"],
                    y0=a["y0"],
                    x1=a["x1"],
                    y1=a["y1"],
                    type=a.get("type"),
                    payload_json=a.get("payload_json"),
                    mapped_paragraph_id=mapped_para_id,
                )
                db.add(ea)
            
            db.commit()
            logger.info("Stored PDF page dimensions and external annotations")
        except Exception as e:
            logger.error(f"Failed to extract/persist PDF annotations: {e}", exc_info=True)
    
    return markdown_content, doc.id


def convert_word_to_markdown(file_path):
    """Convert a Word document to markdown."""
    doc = DocxDocument(file_path)
    markdown_lines = []
    image_counter = 0
    image_dir = Path("data/_assets")
    image_dir.mkdir(exist_ok=True)
    
    # Process all document elements including paragraphs and images
    for element in doc.element.body:
        # Check if it's a paragraph
        if element.tag.endswith('}p'):
            # Get the paragraph from the document by matching
            for para in doc.paragraphs:
                if str(para._element) == str(element):
                    text = para.text.strip()
                    if not text:
                        markdown_lines.append("")
                        continue
                    
                    # Check for heading styles
                    if para.style.name.startswith('Heading 1'):
                        markdown_lines.append(f"# {text}")
                    elif para.style.name.startswith('Heading 2'):
                        markdown_lines.append(f"## {text}")
                    elif para.style.name.startswith('Heading 3'):
                        markdown_lines.append(f"### {text}")
                    else:
                        markdown_lines.append(text)
                    break
        
        # Check if element contains an image
        elif element.tag.endswith('}r'):  # Run element
            for draw in element.iter():
                if draw.tag.endswith('}drawing'):
                    try:
                        # Find the image in the drawing
                        for blip in draw.iter():
                            if blip.tag.endswith('}blip'):
                                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId:
                                    image_part = doc.part.related_parts[rId]
                                    image_bytes = image_part.blob
                                    
                                    # Determine file extension
                                    ext = image_part.content_type.split('/')[-1]
                                    if ext == 'png':
                                        ext = 'png'
                                    elif ext in ['jpeg', 'jpg']:
                                        ext = 'jpg'
                                    else:
                                        ext = 'png'
                                    
                                    # Save image
                                    image_filename = f"{os.path.basename(file_path).replace('.docx', '').replace('.doc', '')}_image_{image_counter}.{ext}"
                                    image_path = image_dir / image_filename
                                    
                                    with open(image_path, "wb") as img_file:
                                        img_file.write(image_bytes)
                                    
                                    # Add image reference to markdown
                                    markdown_lines.append(f"\n![Image](/_assets/{image_filename})\n")
                                    
                                    image_counter += 1
                    except Exception as e:
                        print(f"Error extracting image: {e}")
    
    # Fallback: also check inline shapes (older approach)
    for i, shape in enumerate(doc.inline_shapes):
        if hasattr(shape, 'image') and image_counter == 0:  # Only if no images found yet
            try:
                # Get the image part
                image_part = shape._inline.graphic.graphicData.pic.blip.fill.blip
                rId = image_part.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    related_part = doc.part.related_parts[rId]
                    image_bytes = related_part.blob
                    
                    # Save image
                    image_filename = f"{os.path.basename(file_path).replace('.docx', '').replace('.doc', '')}_image_{image_counter}.png"
                    image_path = image_dir / image_filename
                    
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    # Add image reference to markdown
                    markdown_lines.append(f"\n![Image](/_assets/{image_filename})\n")
                    
                    image_counter += 1
            except Exception as e:
                print(f"Error extracting image from inline shape: {e}")
    
    return "\n\n".join(markdown_lines)


def convert_pdf_to_markdown(file_path):
    """
    Convert a PDF document to markdown with page boundaries and paragraph-to-page mapping.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (markdown_text, list_of_page_breaks, list_of_paragraph_page_mapping)
    """
    from app_server.services.converters.pdf_converter import pdf_converter
    return pdf_converter.extract_text_to_markdown(file_path)
